# -*- coding: utf-8 -*-
"""
文件加载工具：支持读取 .txt / .docx / .pdf / .csv 文件
"""
import os
import pandas as pd


class FileLoader:
    SUPPORTED_EXTENSIONS = (".txt", ".docx", ".pdf", ".md", ".csv")

    @staticmethod
    def read_text(filepath: str) -> str:
        """读取文件为纯文本"""
        ext = os.path.splitext(filepath)[1].lower()
        if ext == ".docx":
            return FileLoader._read_docx(filepath)
        elif ext == ".pdf":
            return FileLoader._read_pdf(filepath)
        elif ext == ".csv":
            return FileLoader._read_csv_as_text(filepath)
        else:
            with open(filepath, "r", encoding="utf-8") as f:
                return f.read()

    @staticmethod
    def read_text_from_bytes(file_bytes, filename: str) -> str:
        """从上传的文件字节读取文本（Streamlit 上传用）"""
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".docx":
            return FileLoader._read_docx_bytes(file_bytes)
        elif ext == ".pdf":
            return FileLoader._read_pdf_bytes(file_bytes)
        else:
            try:
                return file_bytes.decode("utf-8")
            except UnicodeDecodeError:
                return file_bytes.decode("gbk")

    @staticmethod
    def _read_docx(filepath: str) -> str:
        from docx import Document
        doc = Document(filepath)
        return "\n".join([p.text for p in doc.paragraphs])

    @staticmethod
    def _read_docx_bytes(file_bytes) -> str:
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])

    @staticmethod
    def _read_pdf(filepath: str) -> str:
        """读取 PDF 文件，优先 pdfplumber，降级 PyPDF2"""
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            pass
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        except ImportError:
            return "[ERROR] 未安装 pdfplumber 或 PyPDF2，请运行 pip install pdfplumber"

    @staticmethod
    def _read_pdf_bytes(file_bytes) -> str:
        """从字节读取 PDF"""
        import io
        try:
            import pdfplumber
            text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            pass
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        except ImportError:
            return "[ERROR] 未安装 pdfplumber，请运行 pip install pdfplumber"

    @staticmethod
    def _read_csv_as_text(filepath: str) -> str:
        df = pd.read_csv(filepath)
        return df.to_string(index=False)

    @staticmethod
    def read_csv(filepath: str) -> pd.DataFrame:
        return pd.read_csv(filepath)

    @staticmethod
    def list_resumes(resume_dir: str) -> list:
        if not os.path.exists(resume_dir):
            return []
        files = []
        for f in os.listdir(resume_dir):
            if f.lower().endswith(FileLoader.SUPPORTED_EXTENSIONS):
                files.append(os.path.join(resume_dir, f))
        return files

    @staticmethod
    def deduplicate_text(text: str) -> str:
        """
        文本去重：去除重复行、重复段落、多余空白
        用于简历文本预处理，减少 LLM token 消耗
        """
        if not text or not text.strip():
            return text

        import re

        # 1. 统一换行符
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # 2. 按行分割，去除首尾空白
        lines = [line.strip() for line in text.split("\n")]

        # 3. 去除完全重复的行（保留第一次出现的顺序）
        seen_lines = set()
        unique_lines = []
        for line in lines:
            if line == "":
                unique_lines.append(line)
                continue
            line_lower = line.lower()
            if line_lower not in seen_lines:
                seen_lines.add(line_lower)
                unique_lines.append(line)

        # 4. 合并连续空行为单个空行
        result_lines = []
        prev_empty = False
        for line in unique_lines:
            if line == "":
                if not prev_empty:
                    result_lines.append(line)
                prev_empty = True
            else:
                result_lines.append(line)
                prev_empty = False

        # 5. 去除首尾空行
        while result_lines and result_lines[0] == "":
            result_lines.pop(0)
        while result_lines and result_lines[-1] == "":
            result_lines.pop()

        result = "\n".join(result_lines)

        # 6. 去除重复的段落
        paragraphs = re.split(r"\n\s*\n", result)
        seen_paras = set()
        unique_paras = []
        for para in paragraphs:
            para_clean = para.strip().lower()
            if para_clean and para_clean not in seen_paras:
                seen_paras.add(para_clean)
                unique_paras.append(para.strip())

        return "\n\n".join(unique_paras)
